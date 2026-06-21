"""반품 분석 진단·보고서 — 규칙기반(외부 AI 미사용).

진단/액션 로직(diagnose·size_diagnosis·action_priorities·ACTION_PLAN)을 함수로 분리해
화면 표시와 Word(.docx) 보고서가 같은 내용을 쓰도록 함.
입력 df 컬럼: reason, qty, 사이즈, (product_name)
"""
import io

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# 반품 사유 그룹 → (키워드, 진단 문장, 권장 조치)
# ※ 순서 중요: '큼/작음'을 '사이즈(기타)'보다 먼저 둬야 방향이 정확히 분류됨
GROUPS = {
    "사이즈(큼)": (["커요", "크게", "큼"],
               "사이즈가 크게 나온다는 반품이 많습니다.",
               "정사이즈보다 크게 제작됐을 수 있음 — 한 단계 작은 사이즈 안내·실측 보정"),
    "사이즈(작음)": (["작아요", "작게", "작음"],
                "사이즈가 작게 나온다는 반품이 많습니다.",
                "정사이즈보다 작게 제작 — 한 단계 큰 사이즈 안내·실측 보정"),
    "사이즈(기타)": (["사이즈", "크기"],
                "사이즈 선택/확인 관련 반품입니다.",
                "상세페이지 실측 사이즈표·핏 안내 보강"),
    "색상/화면": (["색상", "화면", "사진", "컬러"],
               "실물과 화면 색상 차이로 인한 반품이 많습니다.",
               "자연광 실촬영·색보정 점검, 색상별 실물컷 추가"),
    "소재/품질": (["원단", "소재", "비침", "하자", "불량", "얇", "두께"],
               "소재·품질 관련 반품이 많습니다.",
               "입고 검수 강화, 소재(두께·비침) 정보 상세 표기"),
    "디자인/핏": (["디자인", "착용감", "어울", "핏"],
               "디자인·착용감 관련 반품이 많습니다.",
               "착용컷·디테일컷 보강으로 기대치 정렬"),
    "상세설명": (["상세", "상이", "설명"],
              "상세설명과 실제가 다르다는 반품이 있습니다.",
              "상세페이지 정보 정확성 점검"),
    "단순변심": (["변심", "필요", "잘못", "오주문", "다시", "안어울"],
              "단순 변심 반품입니다.",
              "구매 전 정보 제공 강화(직접 개선은 제한적)"),
}
ACTION_PLAN = [
    ("1. 진단", "반품 사유·사이즈 집중 구간 파악 (본 분석)"),
    ("2. 상세페이지", "실측 사이즈표·핏 안내·실물컷 보강"),
    ("3. 검수", "입고 검수 및 소재(두께·비침) 정보 표기 강화"),
    ("4. CS", "반품 사유 태깅·반복 항목 주간 모니터링"),
    ("5. 재발방지", "상위 사유 개선 후 반품율 추적·검증"),
]
RED, ORANGE, INDIGO = "F8C9C9", "FCE2B5", "4F46E5"


def classify(reason):
    r = str(reason or "")
    for g, (kws, _, _) in GROUPS.items():
        if any(k in r for k in kws):
            return g
    return "기타"


def _group_sums(df):
    g = df.copy()
    g["_g"] = g["reason"].map(classify)
    return g.groupby("_g")["qty"].sum().sort_values(ascending=False)


# ─────────────────────── 진단·액션 (화면·문서 공용) ───────────────────────

def diagnose(df, top=4):
    """[(그룹, 진단문장, 조치, 건수), ...] (건수 많은 순)."""
    out = []
    for g, cnt in _group_sums(df).items():
        if g in GROUPS:
            out.append((g, GROUPS[g][1], GROUPS[g][2], int(cnt)))
    return out[:top]


def action_priorities(df, top=5):
    """[(순위, 그룹, 조치, 건수), ...]."""
    rows = []
    for i, (g, cnt) in enumerate(_group_sums(df).head(top).items(), 1):
        act = GROUPS.get(g, (None, None, "개별 모니터링"))[2]
        rows.append((i, g, act, int(cnt)))
    return rows


def size_diagnosis(df, top=8):
    """[(사이즈, 건수, 주요사유, 조치), ...]."""
    out = []
    for sz, cnt in df.groupby("사이즈")["qty"].sum().sort_values(ascending=False).head(top).items():
        sub = df[df["사이즈"] == sz]
        topr = sub.groupby("reason")["qty"].sum().sort_values(ascending=False)
        reason = topr.index[0] if len(topr) else "-"
        act = GROUPS.get(classify(reason), (None, None, "개별 점검"))[2]
        out.append((str(sz), int(cnt), str(reason), act))
    return out


# ─────────────────────────── Word(.docx) 생성 ───────────────────────────

def _shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def _set(cell, text, *, header=False, shade=None):
    cell.text = str(text)
    p = cell.paragraphs[0]
    run = p.runs[0] if p.runs else p.add_run("")
    run.font.size = Pt(9)
    if header:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(cell, INDIGO)
    elif shade:
        _shade(cell, shade)


def _table(doc, headers):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        _set(t.rows[0].cells[i], h, header=True)
    return t


def _cell_shade(v):
    return RED if v >= 20 else ORANGE if v >= 10 else None


def _save(doc):
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_reason_report(product_name, df):
    total = int(df["qty"].sum()) if not df.empty else 0
    doc = Document()
    doc.add_heading("반품 현황 분석 — 상품별", level=0)
    doc.add_heading(str(product_name), level=1)
    doc.add_paragraph(f"총 반품 건수: {total:,}건 · 반품 사유 {df['reason'].nunique()}종")

    doc.add_heading("반품 사유별", level=2)
    by = (df.groupby("reason", as_index=False).agg(건수=("qty", "sum"))
            .sort_values("건수", ascending=False))
    t = _table(doc, ["반품 사유", "건수", "비율"])
    for _, r in by.iterrows():
        cnt = int(r["건수"])
        c = t.add_row().cells
        _set(c[0], r["reason"])
        _set(c[1], f"{cnt:,}", shade=_cell_shade(cnt))
        _set(c[2], f"{cnt / total * 100:.1f}%" if total else "-")

    doc.add_heading("진단 코멘트", level=2)
    for g, diag, act, cnt in diagnose(df, top=4):
        doc.add_paragraph(f"[{g}] {diag} ({cnt:,}건) → {act}", style="List Bullet")

    doc.add_heading("액션 우선순위", level=2)
    at = _table(doc, ["우선순위", "항목", "권장 조치", "건수"])
    for i, g, act, cnt in action_priorities(df):
        c = at.add_row().cells
        _set(c[0], i); _set(c[1], g); _set(c[2], act); _set(c[3], f"{cnt:,}")
    return _save(doc)


def build_size_report(product_name, df):
    doc = Document()
    doc.add_heading("반품 현황 분석 — 사이즈별", level=0)
    doc.add_heading(str(product_name), level=1)

    doc.add_heading("사이즈 × 사유 교차표", level=2)
    ct = (pd.crosstab(df["사이즈"], df["reason"], values=df["qty"], aggfunc="sum",
                      margins=True, margins_name="합계").fillna(0).astype(int)
          .sort_values("합계", ascending=False))
    cols = list(ct.columns)
    t = _table(doc, ["사이즈"] + [str(c) for c in cols])
    for idx, row in ct.iterrows():
        cells = t.add_row().cells
        _set(cells[0], idx, header=(idx == "합계"))
        for i, c in enumerate(cols, 1):
            v = int(row[c])
            _set(cells[i], f"{v:,}" if v else "-", shade=_cell_shade(v))
    doc.add_paragraph("※ 10건 이상 주황, 20건 이상 빨강으로 강조했습니다.")

    doc.add_heading("사이즈별 진단", level=2)
    for sz, cnt, reason, act in size_diagnosis(df):
        doc.add_paragraph(f"[{sz}] {cnt:,}건 · 주요 사유 '{reason}' → {act}",
                          style="List Bullet")

    doc.add_heading("액션 플랜 (5단계)", level=2)
    at = _table(doc, ["단계", "내용"])
    for s, d in ACTION_PLAN:
        c = at.add_row().cells
        _set(c[0], s); _set(c[1], d)
    return _save(doc)
